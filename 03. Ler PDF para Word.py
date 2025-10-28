import pymupdf

#pip install python-docx
#lib para trabalhar com arquivos Word 

from docx import Document

import os

linhas_texto = []

origem_pdf = r'C:\Rafaélla Sena - Python\Python 2\Python 2 - Clarify\Aluno\Aula_01\bases\Salário Mínimo 2019.pdf'

destino_word = r'C:\Rafaélla Sena - Python\Python 2\Python 2 - Clarify\Aluno\Aula_01\bases\Salário Mínimo 2019.docx'

with pymupdf.open(origem_pdf) as pdf:
    #percorre cada página do PDF
    for pagina in pdf:
        texto = pagina.get_text()

        if texto:
            linhas = texto.split('\n') #quebrar texto por linhas

            linhas_texto.extend(linhas)

#sai do with/FOR

#cria um novo documento Word vazio
documento = Document()

#adicionar titulo no inicio do documento
documento.add_heading('Texto estraido do PDF', level = 1)

#para cada linha de texto extraida do PDF
for linha in linhas_texto:

    #cria uma string vazia para armazenar os caracteres válidos
    linha_limpa = ''
    
    #percorre cada caractere dentro da linha
    #ch = character (caracter)
    for ch in linha:

        #verifica se o caractere é imprimivel 
        if ch.isprintable():

            #se for imprimivel, adicionar caracter a nova linha limpa
            linha_limpa += ch

    #adicionar a linha limpa como um paragrafo no documento Word
    documento.add_paragraph(linha_limpa)

#salvar o documento no WORD
documento.save(destino_word)

#abrir o WORD
os.startfile(destino_word)

























































































